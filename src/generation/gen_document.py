import markdown
from docx import Document
import pandas as pd
from io import StringIO
import re


def split_intro_and_table(response):
    lines = response.strip().splitlines()
    table_start = next((i for i, line in enumerate(lines) if re.match(r"^\s*\|.*\|\s*$", line)), None)

    if table_start is None:
        raise ValueError("No markdown table found.")

    intro = "\n".join(lines[:table_start]).strip()
    table_md = "\n".join(lines[table_start:]).strip()

    return intro, table_md

intro_text, md_table = split_intro_and_table(response)

# Parse the Markdown table using pandas
df = pd.read_csv(StringIO(md_table), sep="|", skipinitialspace=True)
df = df.drop(df.columns[[0, -1]], axis=1)  # Remove border columns

def markdown_table_to_docx(df, intro, filename="table.docx"):
    doc = Document()

    # Add the introduction paragraph
    if intro:
        doc.add_paragraph(intro)

    # Add the table
    table = doc.add_table(rows=1, cols=len(df.columns))
    
    # Add header
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = col_name

    # Add rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    doc.save(f"../../output/filled_tables/{filename}")

markdown_table_to_docx(df, intro_text)
